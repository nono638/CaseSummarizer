"""
File Readers for Non-PDF Document Formats.

Provides text extraction for TXT, RTF, DOCX, and image files (PNG/JPG).
Each reader returns an ExtractionResult with text, confidence, and status.

Supported formats:
    - TXT: Direct file read (UTF-8)
    - RTF: striprtf library conversion
    - DOCX: python-docx paragraph and table extraction
    - PNG/JPG: OCR via Tesseract (delegates to OCRProcessor)

Example usage:
    >>> from src.core.extraction.dictionary_utils import DictionaryTextValidator
    >>> from src.core.extraction.ocr_processor import OCRProcessor
    >>> readers = FileReaders(DictionaryTextValidator(), OCRProcessor(DictionaryTextValidator()))
    >>> result = readers.read_text_file(Path("document.txt"))
    >>> print(result['method'])  # 'direct_read'
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

from PIL import Image

logger = logging.getLogger(__name__)

from .dictionary_utils import DictionaryTextValidator

if TYPE_CHECKING:
    from .extraction_result import ExtractionResult


class FileReaders:
    """
    Reads text from non-PDF document formats.

    Provides a unified interface for extracting text from TXT, RTF, DOCX,
    and image files. Each method returns an ExtractionResult.

    Attributes:
        dictionary: DictionaryTextValidator for confidence calculation
        ocr_processor: Optional OCRProcessor for image files
    """

    def __init__(self, dictionary: DictionaryTextValidator, ocr_processor=None):
        """
        Initialize the file readers.

        Args:
            dictionary: DictionaryTextValidator instance for confidence calculation
            ocr_processor: Optional OCRProcessor for image files
        """
        self.dictionary = dictionary
        self.ocr_processor = ocr_processor

    def _wrap_reader(
        self,
        file_path: Path,
        log_label: str,
        method_name: str,
        error_subject: str,
        read_fn,
    ) -> ExtractionResult:
        """
        Run a reader callback and wrap the result in an ExtractionResult.

        Shared scaffolding for text/RTF/DOCX readers: logs, calls ``read_fn``
        which must return ``(text, page_count)``, computes dictionary
        confidence, and converts exceptions to ExtractionResult errors. If
        ``read_fn`` returns a pre-built ExtractionResult (e.g. an early-out
        error like "no readable text"), it is returned unchanged.

        Args:
            file_path: File being read (used for logging only)
            log_label: Human-readable format name (debug logs)
            method_name: Method tag stored on the ExtractionResult
            error_subject: Subject used in the failure message,
                "Failed to read {error_subject}: ..."
            read_fn: Callable taking no args, returning ``(text, page_count)``
                or an ExtractionResult to short-circuit.

        Returns:
            ExtractionResult.success on text, or ExtractionResult.error on failure.
        """
        from .extraction_result import ExtractionResult

        logger.debug("Reading %s file: %s", log_label, file_path.name)
        try:
            result = read_fn()
            if isinstance(result, ExtractionResult):
                return result
            text, page_count = result
            confidence = self.dictionary.calculate_confidence(text)
            logger.debug("%s dictionary confidence: %.1f%%", log_label, confidence)
            return ExtractionResult.success(text, method_name, confidence, page_count=page_count)
        except Exception as e:
            return ExtractionResult.error(
                f"Failed to read {error_subject}: {e!s}",
                page_count=0,
            )

    def read_text_file(self, file_path: Path) -> ExtractionResult:
        """Read a plain text (.txt) file via UTF-8 direct read."""

        def _read():
            """Read raw UTF-8 text and return (text, 1-page-count)."""
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                return f.read(), 1

        return self._wrap_reader(file_path, "text", "direct_read", "text file", _read)

    def read_rtf_file(self, file_path: Path) -> ExtractionResult:
        """Read a Rich Text Format (.rtf) file via striprtf."""

        def _read():
            """Strip RTF markup and return (plain_text, 1-page-count)."""
            from striprtf.striprtf import rtf_to_text

            with open(file_path, encoding="utf-8", errors="ignore") as f:
                rtf_content = f.read()
            text = rtf_to_text(rtf_content)
            logger.debug("Extracted %d characters from RTF", len(text))
            return text, 1

        return self._wrap_reader(file_path, "RTF", "rtf_extraction", "RTF file", _read)

    def read_docx_file(self, file_path: Path) -> ExtractionResult:
        """Read a Word document (.docx) via python-docx; paragraphs + tables."""
        from .extraction_result import ExtractionResult

        def _read():
            """Pull paragraphs + table cells; short-circuit on empty docs."""
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += "\n" + " | ".join(row_text)
            if not text.strip():
                # Preserve exact original error contract for empty docs.
                return ExtractionResult.error(
                    "Word document contains no readable text.",
                    page_count=0,
                )
            return text, len(doc.sections) or 1

        return self._wrap_reader(
            file_path, "Word document", "docx_extraction", "Word document", _read
        )

    def read_image_file(self, file_path: Path) -> ExtractionResult:
        """
        Read an image file (.png, .jpg, .jpeg) using OCR.

        Delegates to OCRProcessor for Tesseract-based text extraction.

        Args:
            file_path: Path to the image file

        Returns:
            ExtractionResult with text, method='image_ocr', confidence.

        Example:
            >>> readers = FileReaders(DictionaryTextValidator(), OCRProcessor(DictionaryTextValidator()))
            >>> result = readers.read_image_file(Path("scan.png"))
        """
        from .extraction_result import ExtractionResult

        logger.debug("Reading image file: %s", file_path.name)

        if self.ocr_processor is None:
            return ExtractionResult.error(
                "OCR processor not available for image files.",
                page_count=1,
            )

        try:
            with Image.open(file_path) as img:
                result = self.ocr_processor.process_image(img)

            # Add page_count to result
            result.page_count = 1

            return result

        except Exception as e:
            return ExtractionResult.error(
                f"Failed to process image: {e!s}",
                page_count=1,
            )
